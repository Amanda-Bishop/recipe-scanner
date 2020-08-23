# Import statements
from docx import Document
from google.cloud import vision
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
import os, io, sys

# Setting up Cloud Vision
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r"/Users/amanda/Documents/apikey.json"           # Create an api key json file and use it in your program
client = vision.ImageAnnotatorClient()

# Main class
class Window(QMainWindow):
    # Initialization function
    def __init__(self):
        super().__init__()
        self.title = 'Recipe Scanner'
        self.x, self.y, self.w, self.h = 400, 300, 500, 300
        self.doc = Document()
        self.name = 'unnamed document.docx'                                                     # Sets the default name of the file to be 'unnamed document'
        self.selectedFiles = []
        self.initUI()                                                                           # Calls the UI initialization function

    # UI intialization function
    def initUI(self):
        # Sets the window title and size
        self.setWindowTitle(self.title)
        self.setGeometry(self.x, self.y, self.w, self.h)

        # Creates the text on the screen
        self.l1 = QLabel('Add the document\'s name below', self)
        self.l1.move(170,5)
        self.l1.resize(300,30)

        self.l2 = QLabel('If you want to add a picture of the recipe click the button below', self)
        self.l2.move(70, 70)
        self.l2.resize(500, 30)

        self.l3 = QLabel('Press the below button to select a recipe file',self)
        self.l3.move(140,125)
        self.l3.resize(500,30)

        self.l4 = QLabel('You can select multiple files at a time or after already selecting files',self)
        self.l4.move(50, 145)
        self.l4.resize(500, 30)

        self.l5 = QLabel('click the button again to add the new files onto the end of your document',self)
        self.l5.move(40, 165)
        self.l5.resize(500, 30)

        self.l6 = QLabel('',self)
        self.l6.move(250,230)
        self.l6.resize(500,30)

        self.pic = QLabel('',self)
        self.pic.resize(100, 100)
        self.pic.move(320, 60)

        # Creates the text box
        self.t1 = QLineEdit(self)
        self.t1.move(200,35)
        self.t1.resize(120,20)

        # Creates the buttons
        self.b1 = QPushButton('Add Picture',self)
        self.b1.setFont(QFont('Arial', 15))
        self.b1.move(200, 100)
        self.b1.resize(120, 30)
        self.b1.clicked.connect(self.getPic)                                                       # Connects the button to the getPic function

        self.b2 = QPushButton('Get Recipes',self)
        self.b2.setFont(QFont('Arial', 15))
        self.b2.move(200,200)
        self.b2.resize(120,30)
        self.b2.clicked.connect(self.getRecipes)                                                   # Connects the button to the getRecipes function

        self.b3 = QPushButton('Done!',self)
        self.b3.setFont(QFont('Arial', 15))
        self.b3.move(220, 260)
        self.b3.resize(70, 30)
        self.b3.clicked.connect(self.close)                                                        # Exits the program when clicked

        self.show()

    # Function to add a picture to the document
    def getPic(self):
        # Changes the document name if the text box is not empty
        if self.t1.text() != '':
            self.name = self.t1.text() + '.docx'
        self.doc.save(self.name)
        dialog = QFileDialog.getOpenFileName(self, "Get Pictures", "", "Image Files (*.png *.jpg *.bmp)")   # Creates a file selector so the user can pick an image from their computer
        self.doc.add_picture(dialog[0])                                                            # Adds the picture to the document
        pixmap = QPixmap(dialog[0])
        pixmap = pixmap.scaled(25, 25, Qt.KeepAspectRatio)
        self.pic.setPixmap(pixmap)                                                                 # Displays the selected picture on the window
        self.doc.save(self.name)                                                                   # Saves the file

    # Function to add the recipes to the document
    def getRecipes(self):
        # Changes the document name if the text box is not empty
        if self.t1.text() != '':
            self.name = self.t1.text() + '.docx'
        self.doc.save(self.name)
        dialog = QFileDialog.getOpenFileNames(self,"Get Recipes", "", "Image Files (*.png *.jpg *.bmp)")    # Creates a file selector so the user can pick an image of a recipe from their computer
        # Loops through the selected files to add them to the document
        for i in dialog[0]:
            self.addTxt(i)
            file = i.split('/')
            self.selectedFiles.append(file[-1])                                                    # Adds the file name to the selectedFiles array
            self.l6.move(self.l6.x()-50,230)
        self.l6.setText('Selected Files: ' + ', '.join(self.selectedFiles))                        # Displays the selected files

    # Function to add the text to the document
    def addTxt(self, fileName):
        # Opens the file and stores the text detected from the picture
        with io.open(fileName, 'rb') as imageFile:
            content = imageFile.read()
        image = vision.types.Image(content=content)
        response = client.text_detection(image=image)
        for text in response.text_annotations:
            self.doc.add_paragraph(text.description)                                               # Adds the text to the document
            break
        self.doc.save(self.name)                                                                   # Saves the document

# Runs the program
if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = Window()
    sys.exit(app.exec())